import docx
import os
import uuid
from app.converters.base_converter import BaseConverter

class DocxConverter(BaseConverter):
    def convert(self, input_path, target_format):
        doc = docx.Document(input_path)
        output_filename = f"{uuid.uuid4()}.{target_format}"
        output_path = self.get_output_path(output_filename)

        if target_format == "txt":
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write('\n'.join(full_text))

        elif target_format == "html":
            # Very basic DOCX to HTML conversion
            html_content = "<html><body>"
            for para in doc.paragraphs:
                html_content += f"<p>{para.text}</p>"
            html_content += "</body></html>"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)
        
        else:
            raise ValueError(f"Unsupported target format for DOCX: {target_format}")

        return output_filename
