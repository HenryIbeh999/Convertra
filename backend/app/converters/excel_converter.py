import pandas as pd
import docx
import fitz
import os
import uuid
from app.converters.base_converter import BaseConverter

class ExcelConverter(BaseConverter):
    def convert(self, input_path, target_format):
        df = pd.read_excel(input_path)
        output_filename = f"{uuid.uuid4()}.{target_format}"
        output_path = self.get_output_path(output_filename)

        if target_format == "csv":
            df.to_csv(output_path, index=False)
        
        elif target_format == "json":
            df.to_json(output_path, orient='records')

        elif target_format == "docx":
            # Convert Excel to DOCX Table
            doc = docx.Document()
            doc.add_heading('Converted Excel Data', 0)
            
            # Create a table
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr_cells[idx].text = str(col)

            # Add data to the table
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for idx, value in enumerate(row):
                    row_cells[idx].text = str(value)
            
            doc.save(output_path)

        elif target_format == "pdf":
            # Very basic PDF rendering from DataFrame
            # This will create a simple text-based rendering
            doc = fitz.open()
            page = doc.new_page()
            
            # Construct a string representation of the table
            text_repr = df.to_string()
            page.insert_text((50, 50), text_repr, fontsize=10)
            
            doc.save(output_path)
            doc.close()

        else:
            raise ValueError(f"Unsupported target format for Excel: {target_format}")

        return output_filename
